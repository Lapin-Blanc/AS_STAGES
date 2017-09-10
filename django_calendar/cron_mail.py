#import os
#os.environ.setdefault("DJANGO_SETTINGS_MODULE", "AS_STAGES.settings")
#import django
#django.setup()

from io import BytesIO

from django.utils import timezone
from django_calendar.models import Traineeship, Supervisor
from django.core.mail import EmailMessage
from django.conf import settings
from docx import Document
from docx.shared import Pt

JOURS = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']

def send_all_schedules(from_date=timezone.localdate()):

    next_monday = from_date + timezone.timedelta(days=7-from_date.weekday())

    base_body = "Horaires de stages\n---------------------------\n"
    base_subject = "[AS] du lundi %s au dimanche %s" %  (
            next_monday.strftime("%d-%m-%Y"),
            (next_monday + timezone.timedelta(days=6)).strftime("%d-%m-%Y"),
        )

    sender = settings.AS_SENDER
    
    for supervisor in Supervisor.objects.all():
        
        recipients = settings.PERMANENT_RECIPIENTS + [supervisor.email]
        subject = base_subject + " pour %s" % (supervisor)
        body = "Superviseur : %s\n" % (supervisor,)
        msg = EmailMessage(subject, body, sender, recipients)

        # Build and attach documents
        traineeships = Traineeship.objects.filter(is_closed=False, supervisor=supervisor)
        for ts in traineeships:
            student = ts.student
            filename='Horaire %s %s.docx' % (student.last_name, student.first_name)
    
            buffer = BytesIO()
            document = Document()
    
            document.styles["Title"].font.size = Pt(18)
            document.styles["Subtitle"].font.size = Pt(16)
        
            document.add_heading("%s %s : du %s au %s" % (
                ts.student.first_name,
                ts.student.last_name,
                next_monday.strftime("%d-%m-%Y"),
                (next_monday + timezone.timedelta(days=6)).strftime("%d-%m-%Y"),
                )
            ,0)
        
            document.add_paragraph("Stage d'%s - %s" % (ts.category, ts.place,), style="Subtitle")
        
            table = document.add_table(rows=1, cols=5)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Jour'
            hdr_cells[1].text = 'De'
            hdr_cells[2].text = 'A'
            hdr_cells[3].text = 'Périodes'
            hdr_cells[4].text = 'Heures'
            for x in range(7):
                row_day = next_monday + timezone.timedelta(days=x)
                day_periods = ts.periods.filter(start__date=row_day).order_by('start')
                row_cells = table.add_row().cells
                row_cells[0].text = JOURS[x]
                num_p = 0
                for p in day_periods :
                    num_p += 1
                    row_cells[1].text = timezone.localtime(p.start).strftime("%H:%M")
                    row_cells[2].text = timezone.localtime(p.end).strftime("%H:%M")
                    row_cells[3].text = str(p.period_duration())
                    row_cells[4].text = str(p.hour_duration())
                    if not num_p == len(day_periods):
                        row_cells = table.add_row().cells
            document.save(buffer)
             
            body += "%s %s - stage d'%s - %s\n" % (student.last_name, student.first_name, ts.category, ts.place)
            msg.attach(filename, buffer.getvalue())
            buffer.close()

        if traineeships:
            msg.body = body
            msg.send()
