#!/usr/bin/env python
# -*- coding: utf-8 -*-
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.utils import formats, dateparse, timezone
from .models import Period, Traineeship, Student
from django.core.exceptions import ValidationError

from datetime import datetime, date

from io import BytesIO
from docx import Document
from docx.shared import Pt

def json_access_error(request):
     return  JsonResponse(
                {
                  "errors": [
                    {
                      "status": "403",
                      "source": { "pointer": request.path },
                      "detail": "vous n'êtes plus autorisé à utiliser cette période"
                    },
                  ]
                },
                status=403
            )
 
def time_limit():
    today = timezone.localdate()
    days_offset = 3-today.weekday()
    return timezone.make_aware(datetime.combine(today+timezone.timedelta(days=days_offset), datetime.min.time()))

def calendar(request, action, traineeship):
    user = request.user
    traineeship = Traineeship.objects.get(id=int(traineeship))
    try:
        student = user.student
    except Student.DoesNotExist:
        student = None

    # calendar read
    if action=='read':
        time_start = timezone.make_aware(datetime.combine(dateparse.parse_date(request.GET['start']), datetime.min.time()))
        time_end = timezone.make_aware(datetime.combine(dateparse.parse_date(request.GET['end']), datetime.min.time()))
        base_criteria = {
            'traineeship' : traineeship
        }
        if request.GET['type']=='past':
            base_criteria['start__gte'] = time_start
            base_criteria['end__lt'] = time_limit()

        if request.GET['type']=='future':
            base_criteria['start__gte'] = time_limit()
            base_criteria['end__lt'] = time_end
        ps = Period.objects.filter(**base_criteria)
        d = []
        for p in ps:
            d.append({
                    'id': p.id,
                    'start': p.start,
                    'end': p.end,
            })
        return JsonResponse(d, safe=False)

    # create period
    if action=='create':
        time_start = dateparse.parse_datetime(request.GET['start'])
        time_end = dateparse.parse_datetime(request.GET['end'])
        if student and time_start<time_limit():
            return json_access_error(request)            
        try:
            p = traineeship.periods.create(start=time_start, end=time_end)
            return JsonResponse({"event_id" : p.id}, safe=False)
        except ValidationError as e:
            return JsonResponse(
                {
                  "errors": [
                    {
                      "status": "422",
                      "source": { "pointer": request.path },
                      "detail": "%s" % e.args[0]
                    },
                  ]
                },
                status=422
            )
    
    # delete event
    if action=='delete':
        p = traineeship.periods.get(id=int(request.GET['event_id']))
        if student and p.start<time_limit():
            return json_access_error(request)            
        p.delete()
        return JsonResponse({"event_id" : 0}, safe=False)
        
    # update event
    if action=='update':
        try:
            p = traineeship.periods.get(id=int(request.GET['event_id']))
            time_start = dateparse.parse_datetime(request.GET['start'])
            time_end = dateparse.parse_datetime(request.GET['end'])
            if student and time_start<time_limit():
                return json_access_error(request)            
            p.start = time_start
            p.end = time_end
            p.save()
            return JsonResponse({"event_id" : p.id}, safe=False)

        except ValidationError as e:
            return JsonResponse(
                {
                  "errors": [
                    {
                      "status": "422",
                      "source": { "pointer": request.path },
                      "detail": "%s" % e.args[0]
                    },
                  ]
                },
                status=422
            )
    # On ne devrait pas arriver ici...
    return JsonResponse(
        {
            "errors": [
                {
                    "status": "400",
                    "source": { "pointer": request.path },
                    "detail": "action not found"
                },
            ]
        },
        status=400
     )


# DOCX
def download_schedule(request, traineeship):
    user = request.user
    ts = Traineeship.objects.get(id=int(traineeship))
    try:
        student = user.student
    except Student.DoesNotExist:
        student = None

    # Create the HttpResponse object with the appropriate docx headers.
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="horaire.docx"'

    buffer = BytesIO()

    document = Document()
    document.add_heading("%s %s : Stage d'%s" % (ts.student.first_name, ts.student.last_name, ts.category), 0)

    document.save(buffer)
    # Get the value of the BytesIO buffer and write it to the response.
    doc = buffer.getvalue()
    buffer.close()
    response.write(doc)
    return response 

JOURS = ['Lundi', 'Mardi', 'Mercredi', 'Jeudi', 'Vendredi', 'Samedi', 'Dimanche']

def download_schedule_for_student(request, student, from_date=timezone.localdate()):
    next_monday = from_date + timezone.timedelta(days=7-from_date.weekday())
    # télécharge l'horaire d'un étudiant particulier pour la semaine suivant la date fournie ou
    # aujourd'hui si cette date n'est pas fournie
    student = Student.objects.get(id=student)
    #ts = student.traineeships.filter(date_start__lte=from_date, is_closed=False)[0]
    ts = student.traineeships.filter(is_closed=False)[0]
    # TODO : pas de stage ouvert, plus d'un stage ouvert, étudiant n'existant pas

    # Create the HttpResponse object with the appropriate docx headers.
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="horaire %s %s.docx"' % (student.last_name,
            student.first_name)

    buffer = BytesIO()

    document = Document()
    document.styles["Title"].font.size = Pt(18)
    document.styles["Subtitle"].font.size = Pt(16)

    document.add_heading("%s %s : du %s au %s" % (
        ts.student.first_name,
        ts.student.last_name,
        next_monday.strftime("%d-%m-%Y"),
        (next_monday + timezone.timedelta(days=6)).strftime("%d-%m-%Y"),
        )
    ,0)

    document.add_paragraph("Stage d'%s - %s" % (ts.category, ts.place,), style="Subtitle")

    table = document.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Jour'
    hdr_cells[1].text = 'De'
    hdr_cells[2].text = 'A'
    hdr_cells[3].text = 'Périodes'
    hdr_cells[4].text = 'Heures'
    for x in range(7):
        row_day = next_monday + timezone.timedelta(days=x)
        day_periods = ts.periods.filter(start__date=row_day).order_by('start')
        row_cells = table.add_row().cells
        row_cells[0].text = JOURS[x]
        num_p = 0
        for p in day_periods :
            num_p += 1
            row_cells[1].text = timezone.localtime(p.start).strftime("%H:%M")
            row_cells[2].text = timezone.localtime(p.end).strftime("%H:%M")
            row_cells[3].text = str(p.period_duration())
            row_cells[4].text = str(p.hour_duration())
            if not num_p == len(day_periods):
                row_cells = table.add_row().cells
    document.save(buffer)
    # Get the value of the BytesIO buffer and write it to the response.
    doc = buffer.getvalue()
    buffer.close()
    response.write(doc)
    return response 

   
